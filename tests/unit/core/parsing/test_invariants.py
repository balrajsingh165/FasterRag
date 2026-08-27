"""Parsing-layer invariants, generated rather than hand-picked.

Every chunk offset, page number, and section label a citation carries is copied out of a
block, so a parser that gets an offset wrong corrupts the whole chain quietly: retrieval
still returns something, it just no longer points at the text it claims. The example suite
in ``test_parsers.py`` checks these on documents we wrote; this one checks them on
documents we did not.

Eight defects came out of writing it, all of them true in general and false only on the
hand-picked inputs the example suite carried. Each was reproduced on ``main`` first and
each is mutation-verified: reverting any one of the eight fixes turns this file red.

Correctness of the text that reaches the index:

* A UTF-8 BOM survived ``decode`` because U+FEFF is not whitespace, so ``str.strip`` kept
  it. A BOM'd Markdown file lost every heading (``<BOM># Title`` is a paragraph), a BOM'd
  CSV named its first column ``<BOM>name``, and BOM'd JSON failed to parse outright. The
  lenient fallback is a second ``decode`` call and needed the same fix.
* A whitespace-only CSV cell rendered as ``column:   `` — the builder strips the block as a
  whole, not each pair, so a file of spaces and commas reached the index as punctuation.
* A JSON scalar with no content rendered as ``": None"`` for ``null`` and a bare ``":"``
  for ``""``: a chunk that gets embedded and retrieved while saying nothing.
* A ``Heading 0`` style gave a docx block outline depth 0, which outranks every real
  heading in the builder's stack and nests every h1 section underneath it.

Errors that escaped the taxonomy, each reaching the API as a generic 500 with no
problem+json body:

* ``csv.Sniffer`` can return a dialect whose delimiter equals its quote character, which
  ``csv.reader`` rejects with a plain ``ValueError`` — not even a ``csv.Error``.
* ``io.StringIO(text)`` defaults to translating newlines, so a lone CR made the csv module
  raise ``new-line character seen in unquoted field`` on a file every spreadsheet opens.
* Any remaining ``csv.Error`` from the row loop escaped untyped; it is now a ``ParseError``.

And one resource bug:

* ``parse_document`` read the whole file and then compared its length to ``max_bytes``, so
  the limit that exists to keep an oversized document out of memory admitted it to memory
  first. The size now comes from ``stat`` and the file is never opened.

CRITICAL: the alphabets below deliberately include a BOM, CR, CRLF, NUL and other C0
controls, CJK, RTL, combining marks, zero-width and line/paragraph separators. Narrowing
any of them to "printable ASCII" puts most of the defects above out of reach of every
generator here while leaving the suite green. Two of the eight — the ``Heading 0`` style
and a BOM in front of undecodable bytes — are out of reach of generation regardless and are
pinned as constructed cases; both survived a mutation run before those were written.
"""

import io
import json
from collections.abc import Callable, Sequence
from itertools import pairwise
from pathlib import Path
from typing import Any, Final, Literal

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from hypothesis import HealthCheck, assume, example, given, settings
from hypothesis import strategies as st
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from fasterrag.core.parsing import (
    parse_bytes,
    parse_csv,
    parse_document,
    parse_docx,
    parse_html,
    parse_json,
    parse_markdown,
    parse_pdf,
    parse_plaintext,
)
from fasterrag.core.parsing.models import (
    BLOCK_SEPARATOR,
    Block,
    DocumentBuilder,
    ParsedDocument,
    ParseFlag,
)
from fasterrag.errors import ErrorCode, FasterRagError, IngestionError, ParseError

BOM = "﻿"
MAX_HEADING_LEVEL = 6

Parser = Callable[[bytes], ParsedDocument]

TEXT_PARSERS: list[Parser] = [parse_plaintext, parse_markdown, parse_csv, parse_html]
PARSER_IDS = ["plaintext", "markdown", "csv", "html"]
EVERY_TEXT_PARSER = pytest.mark.parametrize("parse", TEXT_PARSERS, ids=PARSER_IDS)

# Anything that is not a surrogate: control characters, unassigned code points, and the
# whole astral range are all things a real corpus contains and a parser has to survive.
_SURROGATES: Final[Sequence[Literal["Cs"]]] = ("Cs",)
_ANY_CHARACTER = st.characters(blacklist_categories=_SURROGATES)

# CRITICAL: sampled explicitly because uniform character generation reaches none of these
# at a useful rate, and each one is a shape that has broken a parser here or elsewhere.
_LOADED_FRAGMENTS = [
    BOM,
    "\r\n",
    "\r",
    "\n",
    "\n\n",
    "\t",
    "\x00",
    "\x0b",
    "\x0c",
    "\x1c",
    "\x85",
    "\u2028",
    "\u200b",
    "\xa0",
    "é",
    "é",
    "日本語",
    "العربية",
    "🙂",
    "# ",
    "## ",
    "####### ",
    "- ",
    "> ",
    "```",
    "| a | b |",
    "| --- |",
    '"',
    ",",
    ";",
    "<p>",
    "</p>",
    "<h1>",
    "</h1>",
    "text",
]

_FRAGMENT = st.one_of(
    st.sampled_from(_LOADED_FRAGMENTS),
    st.text(alphabet=_ANY_CHARACTER, min_size=1, max_size=10),
)

SOURCE = st.lists(_FRAGMENT, max_size=20).map("".join)

# The CRLF comparison rewrites every LF into a CRLF, which is only a faithful rewrite of a
# source that has no CR of its own.
LF_SOURCE = SOURCE.filter(lambda source: "\r" not in source)

# CRITICAL: everything here is removed by ``str.strip``, plus an optional leading BOM, so a
# source built from it carries no content at all. Two things are deliberately absent: NUL,
# which Python does not strip, so a document of NUL bytes has content and keeps its block;
# and a BOM anywhere but the front, which is a zero-width no-break space and is content.
BLANK_SOURCE = st.tuples(
    st.sampled_from(["", BOM]),
    st.lists(
        st.sampled_from([" ", "\t", "\n", "\r", "\r\n", "\x0b", "\x0c", "\x1c", "\x85"]),
        max_size=12,
    ).map("".join),
).map("".join)

JSON_VALUE = st.recursive(
    st.one_of(
        st.none(),
        st.booleans(),
        st.integers(),
        st.text(alphabet=_ANY_CHARACTER, max_size=12),
    ),
    lambda inner: st.one_of(
        st.lists(inner, max_size=4),
        st.dictionaries(st.text(alphabet=_ANY_CHARACTER, max_size=8), inner, max_size=4),
    ),
    max_leaves=12,
)

SETTINGS = settings(
    max_examples=60,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture],
)


def assert_offsets_index_the_text(document: ParsedDocument) -> None:
    """Every block names the text it actually sits on."""
    for block in document.blocks:
        assert 0 <= block.start < block.end <= len(document.text)
        assert document.text[block.start : block.end] == block.text


def assert_blocks_tile_the_text(document: ParsedDocument) -> None:
    """Blocks cover the text in order, parted by exactly one separator and nothing else."""
    assert BLOCK_SEPARATOR.join(block.text for block in document.blocks) == document.text

    if document.blocks:
        assert document.blocks[0].start == 0
        assert document.blocks[-1].end == len(document.text)

    for earlier, later in pairwise(document.blocks):
        assert later.start == earlier.end + len(BLOCK_SEPARATOR)


def assert_blocks_carry_content(document: ParsedDocument) -> None:
    """No block is empty or padded; an empty chunk embeds to a vector that means nothing."""
    for block in document.blocks:
        assert block.text
        assert block.text == block.text.strip()


def assert_lookups_agree_with_the_blocks(document: ParsedDocument) -> None:
    """``section_at`` and ``page_at`` answer for the block an offset is inside."""
    for offset in range(len(document.text)):
        block = document.block_at(offset)
        if block is None:
            continue
        assert document.section_at(offset) == block.section
        assert document.page_at(offset) == block.page


def assert_heading_levels_are_in_range(document: ParsedDocument) -> None:
    """Outline depth drives the heading stack, so out-of-range depth reorders sections."""
    for block in document.blocks:
        if block.level is not None:
            assert 1 <= block.level <= MAX_HEADING_LEVEL


def assert_sections_follow_the_headings(document: ParsedDocument) -> None:
    """A block's section is the path of the nearest heading at or above it."""
    nearest: Block | None = None
    for block in document.blocks:
        if block.kind == "heading":
            nearest = block
        assert block.section == (None if nearest is None else nearest.section)


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
def test_block_offsets_index_the_document_text(parse: Parser, source: str) -> None:
    """The one invariant every citation in the system is built on."""
    assert_offsets_index_the_text(parse(source.encode()))


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
def test_blocks_tile_the_text_without_gaps_or_overlap(parse: Parser, source: str) -> None:
    assert_blocks_tile_the_text(parse(source.encode()))


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
def test_every_block_carries_content(parse: Parser, source: str) -> None:
    assert_blocks_carry_content(parse(source.encode()))


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
def test_a_source_with_no_text_yields_no_blocks(parse: Parser, source: str) -> None:
    """A single empty block would index a meaningless vector and cite an empty span."""
    document = parse(source.encode())

    assert bool(document.blocks) == bool(document.text.strip())


@EVERY_TEXT_PARSER
@given(source=BLANK_SOURCE)
@SETTINGS
@example(source="")
@example(source="  \t \r\n \r\n")
@example(source=BOM)
def test_a_blank_source_yields_an_empty_document(parse: Parser, source: str) -> None:
    """The generated form of the case above, and the one that found the CSV defect.

    A file of nothing but spaces parsed to a single block reading ``":"`` — the header and
    the row were both blank, and only the values ``None`` and ``""`` were being dropped.
    """
    document = parse(source.encode())

    assert document.blocks == ()
    assert document.text == ""


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
def test_the_lookups_agree_with_the_block_at_every_offset(parse: Parser, source: str) -> None:
    assert_lookups_agree_with_the_blocks(parse(source.encode()))


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
def test_the_lookups_are_total_over_the_text(parse: Parser, source: str) -> None:
    """Chunk boundaries land on arbitrary offsets, including past the end of the text."""
    document = parse(source.encode())

    for offset in (-1, 0, len(document.text) // 2, len(document.text), len(document.text) + 8):
        document.block_at(offset)
        document.section_at(offset)
        document.page_at(offset)


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
def test_block_at_finds_everything_outside_the_separators(parse: Parser, source: str) -> None:
    """The offsets no block owns are exactly the separators the builder inserted."""
    document = parse(source.encode())
    covered = {offset for block in document.blocks for offset in range(block.start, block.end)}

    for offset in range(len(document.text)):
        assert (document.block_at(offset) is not None) == (offset in covered)

    assert len(document.text) - len(covered) == len(BLOCK_SEPARATOR) * max(
        len(document.blocks) - 1, 0
    )


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
def test_sections_follow_the_heading_path(parse: Parser, source: str) -> None:
    assert_sections_follow_the_headings(parse(source.encode()))


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
def test_heading_levels_stay_within_the_documented_range(parse: Parser, source: str) -> None:
    assert_heading_levels_are_in_range(parse(source.encode()))


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
def test_parsing_the_same_bytes_twice_gives_the_same_document(parse: Parser, source: str) -> None:
    """Content hashes, golden files, and incremental re-indexing all assume this."""
    data = source.encode()

    assert parse(data) == parse(data)


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
@example(source="# Title\n\nBody")
@example(source="name,year\nAlice,2024\n")
def test_a_leading_byte_order_mark_changes_nothing(parse: Parser, source: str) -> None:
    """The defect: U+FEFF is not whitespace, so stripping a block never removed it."""
    assume(not source.startswith(BOM))

    assert parse((BOM + source).encode()) == parse(source.encode())


@EVERY_TEXT_PARSER
@given(source=LF_SOURCE)
@SETTINGS
# The input that caught csv.Sniffer disagreeing with itself across line endings on
# Python 3.12: a space delimiter for the LF form, no delimiter at all for the CRLF one.
# Pinned because generation found it once and will not reliably find it again --
# reverting the normalisation left this file green without it.
@example(source="# \n\n# ")
def test_crlf_and_lf_sources_differ_only_by_the_carriage_returns(
    parse: Parser, source: str
) -> None:
    """Windows line endings must not change what a block is, only what is inside one.

    Only the carriage returns are allowed to differ: markdown-it and the csv module
    normalise them away entirely, while lxml keeps a CR inside a text node, so the shared
    property is the weaker one.
    """
    lf = parse(source.encode())
    crlf = parse(source.replace("\n", "\r\n").encode())

    assert [(block.kind, block.text.replace("\r", "")) for block in crlf.blocks] == [
        (block.kind, block.text) for block in lf.blocks
    ]


@EVERY_TEXT_PARSER
@given(source=SOURCE)
@SETTINGS
@example(source='\x81"\x94ꦼî')
@example(source="«Ë\x94\r\xd7ú칳")
def test_no_parser_lets_an_untyped_error_escape(parse: Parser, source: str) -> None:
    """Every parser failure has to arrive as a ``FasterRagError``.

    Both pinned examples crashed ``parse_csv``: a sniffed delimiter equal to the quote
    character raised ``ValueError``, and a lone CR raised ``csv.Error``. Neither is in the
    taxonomy, so a worker would have died on a document instead of dead-lettering it.
    """
    try:
        parse(source.encode())
    except FasterRagError:
        return


@given(data=st.binary(max_size=200))
@settings(max_examples=100, deadline=None)
def test_arbitrary_bytes_parse_without_an_untyped_error(data: bytes) -> None:
    for parse in TEXT_PARSERS:
        try:
            parse(data)
        except FasterRagError:
            continue


@given(data=st.binary(max_size=120))
@settings(max_examples=150, deadline=None)
def test_undecodable_bytes_are_flagged_rather_than_failing(data: bytes) -> None:
    """A mangled encoding must be visible in the flags, never a silent partial parse."""
    try:
        data.decode("utf-8")
    except UnicodeDecodeError:
        decodable = False
    else:
        decodable = True

    flags = parse_plaintext(data).flags

    assert (ParseFlag.ENCODING_FALLBACK.value in flags) is not decodable


@given(tail=st.binary(max_size=60))
@settings(max_examples=80, deadline=None)
def test_the_lenient_decode_strips_the_mark_too(tail: bytes) -> None:
    """A file can be both BOM'd and mangled, and the fallback is a separate ``decode`` call.

    Fixing only the strict path leaves the mark in place on exactly the documents that are
    already damaged — and those are the ones whose flags say to look at them, so the stray
    U+FEFF lands in the block a human is about to read.
    """
    data = b"\xef\xbb\xbf" + tail + b"\xff\xfe"
    assume(b"".join([tail, b"\xff\xfe"]))

    document = parse_plaintext(data)

    assert ParseFlag.ENCODING_FALLBACK.value in document.flags
    assert not document.text.startswith("﻿")


@given(source=SOURCE)
@SETTINGS
def test_plaintext_loses_no_character_of_its_source(source: str) -> None:
    """Only whitespace is ever removed: everything else has to come out the other side.

    This is the test that non-ASCII content survives. A decode that mangled CJK, dropped a
    combining mark, or split an astral pair would change the character sequence here even
    though every offset invariant still held.
    """
    assume(not source.startswith(BOM))
    document = parse_plaintext(source.encode())

    assert [c for c in document.text if not c.isspace()] == [c for c in source if not c.isspace()]


def _json_bytes(value: Any) -> bytes:
    """Serialize a generated value the way a caller's file would hold it."""
    return json.dumps(value, ensure_ascii=False).encode()


@given(value=JSON_VALUE)
@settings(max_examples=80, deadline=None)
def test_json_holds_the_offset_invariants(value: Any) -> None:
    document = parse_json(_json_bytes(value))

    assert_offsets_index_the_text(document)
    assert_blocks_tile_the_text(document)
    assert_blocks_carry_content(document)
    assert_lookups_agree_with_the_blocks(document)
    assert_heading_levels_are_in_range(document)
    assert_sections_follow_the_headings(document)


@given(value=JSON_VALUE)
@settings(max_examples=80, deadline=None)
def test_json_survives_a_byte_order_mark(value: Any) -> None:
    """A BOM'd JSON file was a hard ``ParseError``: ``json.loads`` rejects U+FEFF."""
    data = _json_bytes(value)

    assert parse_json(BOM.encode() + data) == parse_json(data)


@pytest.mark.parametrize("empty", [b"{}", b"[]", b"null", b'""', b'{"a": null, "b": ""}'])
def test_json_without_content_yields_no_blocks(empty: bytes) -> None:
    """``null`` used to produce a block reading ``": None"`` — an empty block in disguise."""
    document = parse_json(empty)

    assert document.blocks == ()
    assert document.text == ""


# CRITICAL: curated rather than generated because XML forbids the C0 controls outright, so
# a docx cannot contain them and a generator that emitted them would fail in lxml rather
# than in the parser. Everything XML does permit and this parser might mishandle is here:
# BOM, tab, newline, CJK, RTL, combining marks, zero-width, and an astral pair.
_DOCX_FRAGMENTS = ["a", "Z", "0", " ", "\t", "\n", BOM, "é", "é", "日本語", "العربية", "🙂"]
DOCX_TEXT = (
    st.lists(st.sampled_from(_DOCX_FRAGMENTS), min_size=1, max_size=8)
    .map("".join)
    .map(str.strip)
    .filter(bool)
)

DOCX_ITEM = st.one_of(
    st.tuples(st.just("heading"), DOCX_TEXT, st.integers(min_value=0, max_value=9)),
    st.tuples(st.just("paragraph"), DOCX_TEXT, st.just(0)),
    st.tuples(st.just("list_item"), DOCX_TEXT, st.just(0)),
    st.tuples(st.just("table"), DOCX_TEXT, st.just(0)),
)


def _docx_bytes(items: Sequence[tuple[str, str, int]]) -> bytes:
    """Build a docx whose body children follow ``items`` in order."""
    document = Document()
    for kind, text, level in items:
        if kind == "heading":
            document.add_heading(text, level=min(level, MAX_HEADING_LEVEL))
        elif kind == "list_item":
            document.add_paragraph(text, style="List Bullet")
        elif kind == "table":
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = text
            table.cell(0, 1).text = "cell"
        else:
            document.add_paragraph(text)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@given(items=st.lists(DOCX_ITEM, max_size=8))
@settings(max_examples=30, deadline=None)
def test_docx_holds_the_offset_invariants(items: Sequence[tuple[str, str, int]]) -> None:
    document = parse_docx(_docx_bytes(items))

    assert_offsets_index_the_text(document)
    assert_blocks_tile_the_text(document)
    assert_blocks_carry_content(document)
    assert_lookups_agree_with_the_blocks(document)
    assert_heading_levels_are_in_range(document)
    assert_sections_follow_the_headings(document)


@given(items=st.lists(DOCX_ITEM, max_size=8))
@settings(max_examples=30, deadline=None)
def test_docx_keeps_body_order(items: Sequence[tuple[str, str, int]]) -> None:
    """Blocks come back in body order, tables included.

    Iterating paragraphs and tables separately would pass every offset check above and
    still hand chunking a document whose tables had migrated to the end.
    """
    document = parse_docx(_docx_bytes(items))
    expected = ["heading" if kind == "heading" else kind for kind, _, _ in items]

    assert [block.kind for block in document.blocks] == expected
    for (kind, text, _), block in zip(items, document.blocks, strict=True):
        if kind != "table":
            assert block.text == text
        else:
            assert block.text.startswith(text.replace("\n", " "))


@given(items=st.lists(DOCX_ITEM, min_size=1, max_size=6))
@settings(max_examples=25, deadline=None)
def test_docx_parsing_is_deterministic(items: Sequence[tuple[str, str, int]]) -> None:
    data = _docx_bytes(items)

    assert parse_docx(data) == parse_docx(data)


def test_a_heading_zero_style_does_not_outrank_every_real_heading() -> None:
    """Constructed, because ``_docx_bytes`` clamps the level it is given and never emits it.

    Word itself stops at ``Heading 1``, but a converter can define any style it likes, and
    ``Heading 0`` parsed to outline depth 0. The builder keeps only entries strictly
    shallower than the incoming level, and nothing is shallower than 0, so the entry never
    left the stack: every later ``Heading 1`` nested underneath it and the whole document
    came back sectioned under one bogus root.
    """
    document = Document()
    document.styles.add_style("Heading 0", WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph("Front Matter", style="Heading 0")
    document.add_heading("Chapter One", level=1)
    document.add_paragraph("Body of chapter one.")

    buffer = io.BytesIO()
    document.save(buffer)
    parsed = parse_docx(buffer.getvalue())

    assert [block.level for block in parsed.blocks if block.kind == "heading"] == [1, 1]
    assert parsed.blocks[-1].section == "Chapter One"


PDF_LINE = st.tuples(
    st.text(
        alphabet=st.characters(min_codepoint=32, max_codepoint=126),
        min_size=1,
        max_size=40,
    ).filter(lambda line: line.strip() != ""),
    st.sampled_from([8.0, 10.0, 12.0, 18.0, 24.0]),
)
PDF_PAGE = st.lists(PDF_LINE, min_size=1, max_size=6)


def _pdf_bytes(pages: Sequence[Sequence[tuple[str, float]]]) -> bytes:
    """Build a PDF with the given lines and type sizes, one list per page."""
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    for index, lines in enumerate(pages):
        if index:
            pdf.showPage()
        for row, (text, size) in enumerate(lines):
            pdf.setFont("Helvetica", size)
            pdf.drawString(72, 720 - row * 40, text)
    pdf.save()
    return buffer.getvalue()


@given(pages=st.lists(PDF_PAGE, min_size=1, max_size=3))
@settings(max_examples=20, deadline=None)
def test_pdf_holds_the_offset_invariants(pages: Sequence[Sequence[tuple[str, float]]]) -> None:
    """The heading heuristic reorders nothing: it only relabels."""
    try:
        document = parse_pdf(_pdf_bytes(pages))
    except ParseError:
        return

    assert_offsets_index_the_text(document)
    assert_blocks_tile_the_text(document)
    assert_blocks_carry_content(document)
    assert_lookups_agree_with_the_blocks(document)
    assert_heading_levels_are_in_range(document)
    assert_sections_follow_the_headings(document)


@given(pages=st.lists(PDF_PAGE, min_size=1, max_size=3))
@settings(max_examples=20, deadline=None)
def test_pdf_pages_never_go_backwards(pages: Sequence[Sequence[tuple[str, float]]]) -> None:
    """``page_at`` is what a citation prints; a page that moves backwards misattributes."""
    try:
        document = parse_pdf(_pdf_bytes(pages))
    except ParseError:
        return

    numbers = [block.page for block in document.blocks]
    assert all(number is not None and 1 <= number <= len(pages) for number in numbers)
    assert numbers == sorted(number for number in numbers if number is not None)

    seen = [document.page_at(offset) for offset in range(len(document.text))]
    assert seen == sorted(page for page in seen if page is not None)


PAGED_BLOCK = st.tuples(
    st.text(alphabet=_ANY_CHARACTER, min_size=1, max_size=12).filter(lambda t: t.strip() != ""),
    st.integers(min_value=0, max_value=3),
)


@given(entries=st.lists(PAGED_BLOCK, max_size=12))
@settings(max_examples=80, deadline=None)
def test_page_lookup_tracks_a_paginated_document(entries: Sequence[tuple[str, int]]) -> None:
    """Driven through the builder so pages can be generated without building PDFs."""
    builder = DocumentBuilder(mime_type="application/pdf", parser="test")
    page = 1
    for text, advance in entries:
        page += advance
        builder.add("paragraph", text, page=page)
    document = builder.build()

    assert_offsets_index_the_text(document)
    assert_blocks_tile_the_text(document)

    seen = [document.page_at(offset) for offset in range(len(document.text))]
    assert seen == sorted(number for number in seen if number is not None)
    for offset in range(len(document.text)):
        block = document.block_at(offset)
        if block is not None:
            assert document.page_at(offset) == block.page


OVERSIZE = 64 * 1024


@given(limit=st.integers(min_value=1, max_value=1024))
@settings(
    max_examples=25, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture]
)
def test_an_oversized_file_is_refused_before_it_is_read(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, limit: int
) -> None:
    """The limit exists to keep a document out of memory, so reading first defeats it.

    ``workers.cpu_pool._read`` had always taken the size from the directory entry;
    ``parse_document`` read the bytes and measured them afterwards, which admits a file
    larger than memory before refusing it.
    """
    path = tmp_path / "big.txt"
    path.write_bytes(b"x" * OVERSIZE)

    def forbidden(self: Path) -> bytes:
        raise AssertionError(f"{self} was read before its size was checked")

    monkeypatch.setattr(Path, "read_bytes", forbidden)

    with pytest.raises(IngestionError) as caught:
        parse_document(path, max_bytes=limit)

    assert caught.value.code is ErrorCode.PAYLOAD_TOO_LARGE


@given(size=st.integers(min_value=1, max_value=512), slack=st.integers(min_value=0, max_value=64))
@settings(max_examples=60, deadline=None)
def test_the_limit_is_a_ceiling_and_not_a_bound(size: int, slack: int) -> None:
    data = b"x" * size

    assert parse_bytes(data, filename="ok.txt", max_bytes=size + slack).text == data.decode()

    with pytest.raises(IngestionError):
        parse_bytes(data, filename="ok.txt", max_bytes=size - 1)


@given(size=st.integers(min_value=1, max_value=256))
@settings(max_examples=40, deadline=None)
def test_the_limit_is_checked_before_the_type_is_dispatched(size: int) -> None:
    """An unsupported suffix would raise ``ParseError``; oversize has to win."""
    with pytest.raises(IngestionError) as caught:
        parse_bytes(b"x" * size, filename="archive.zip", max_bytes=size - 1)

    assert caught.value.code is ErrorCode.PAYLOAD_TOO_LARGE


@given(size=st.integers(min_value=0, max_value=256))
@settings(
    max_examples=40, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture]
)
def test_a_file_within_the_limit_is_read_normally(tmp_path: Path, size: int) -> None:
    path = tmp_path / "ok.txt"
    path.write_bytes(b"x" * size)

    document = parse_document(path, max_bytes=size)

    assert document.text == "x" * size
