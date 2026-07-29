"""Fixture builders for the parser suites.

DOCX and PDF inputs are generated in-process rather than committed as binaries, so the
inputs stay reviewable and deterministic. What *is* committed is the golden output
snapshot: any change to parser behavior shows up as a reviewed diff
(``docs/failure-modes.md`` row 3).
"""

import io
import json
import os
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from fasterrag.core.parsing.models import ParsedDocument

GOLDEN_DIR = Path(__file__).parent / "golden"
UPDATE_GOLDEN_ENV = "FASTERRAG_UPDATE_GOLDEN"
FIXED_TIMESTAMP = datetime(2026, 1, 1, tzinfo=UTC)

MARKDOWN = """\
# Vendor Agreement

Intro paragraph about the agreement.

## 3. Termination

### 3.2 Notice

Either party may terminate with 30 days written notice.

- first obligation
- second obligation

| Term | Days |
| --- | --- |
| Notice | 30 |
| Cure | 15 |

```python
print("hello")
```
"""

HTML = """\
<html>
  <head>
    <title>Vendor Agreement</title>
    <meta name="author" content="Legal Team">
  </head>
  <body>
    <nav>skip this navigation</nav>
    <script>console.log("skip this");</script>
    <h1>Vendor Agreement</h1>
    <p>Intro paragraph about the agreement.</p>
    <h2>3. Termination</h2>
    <p>Either party may terminate with 30 days written notice.</p>
    <ul><li>first obligation</li><li>second obligation</li></ul>
    <table>
      <tr><th>Term</th><th>Days</th></tr>
      <tr><td>Notice</td><td>30</td></tr>
    </table>
    <footer>skip this footer</footer>
  </body>
</html>
"""

CSV = "name,department,year\nAlice,legal,2024\nBob,finance,2020\n"

JSON_SOURCE = json.dumps(
    {
        "agreement": {
            "title": "Vendor Agreement",
            "termination": {"notice_days": 30, "cure_days": 15},
            "parties": ["Acme", "Globex"],
        }
    },
    indent=2,
)


def docx_bytes() -> bytes:
    """Build a DOCX with headings, a list, a table, and text after the table."""
    document = Document()
    document.core_properties.title = "Vendor Agreement"
    document.core_properties.author = "Legal Team"
    document.core_properties.created = FIXED_TIMESTAMP
    document.core_properties.modified = FIXED_TIMESTAMP

    document.add_heading("Vendor Agreement", level=1)
    document.add_paragraph("Intro paragraph about the agreement.")
    document.add_heading("3. Termination", level=2)
    document.add_paragraph("Either party may terminate with 30 days written notice.")
    document.add_paragraph("first obligation", style="List Bullet")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Term"
    table.cell(0, 1).text = "Days"
    table.cell(1, 0).text = "Notice"
    table.cell(1, 1).text = "30"

    document.add_paragraph("Text that follows the table.")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def pdf_bytes(*, pages: int = 1, blank_second_page: bool = False) -> bytes:
    """Build a PDF whose larger type marks headings, plus body text."""
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    pdf.setTitle("Vendor Agreement")
    pdf.setAuthor("Legal Team")

    pdf.setFont("Helvetica-Bold", 20)
    pdf.drawString(72, 720, "Vendor Agreement")
    pdf.setFont("Helvetica", 10)
    pdf.drawString(72, 690, "Intro paragraph about the agreement.")
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(72, 660, "3. Termination")
    pdf.setFont("Helvetica", 10)
    pdf.drawString(72, 630, "Either party may terminate with 30 days written notice.")

    for _ in range(pages - 1):
        pdf.showPage()
        if not blank_second_page:
            pdf.setFont("Helvetica", 10)
            pdf.drawString(72, 720, "Second page body text continues here.")

    pdf.save()
    return buffer.getvalue()


def blank_pdf_bytes() -> bytes:
    """Build a PDF with a page but no text, standing in for an unreadable scan."""
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def snapshot(document: ParsedDocument) -> dict[str, Any]:
    """Reduce a parsed document to the stable shape golden files record."""
    return {
        "parser": document.parser,
        "mime_type": document.mime_type,
        "flags": list(document.flags),
        "metadata": {
            key: document.metadata[key]
            for key in sorted(document.metadata)
            if key not in {"created", "modified"}
        },
        "blocks": [
            {
                "kind": block.kind,
                "text": block.text,
                "page": block.page,
                "section": block.section,
                "level": block.level,
            }
            for block in document.blocks
        ],
    }


def assert_matches_golden(name: str, document: ParsedDocument) -> None:
    """Compare a parse against its committed golden file.

    Set ``FASTERRAG_UPDATE_GOLDEN=1`` to rewrite the goldens after an intentional parser
    change; the resulting diff is reviewed like code.
    """
    GOLDEN_DIR.mkdir(exist_ok=True)
    path = GOLDEN_DIR / f"{name}.json"
    actual = snapshot(document)

    if os.environ.get(UPDATE_GOLDEN_ENV):
        path.write_text(json.dumps(actual, indent=2) + "\n", encoding="utf-8")
        return

    if not path.exists():
        pytest.fail(f"missing golden file {path.name}; regenerate with {UPDATE_GOLDEN_ENV}=1")

    expected = json.loads(path.read_text(encoding="utf-8"))
    assert actual == expected


@pytest.fixture
def sources(tmp_path: Path) -> Path:
    """Write one file of every text-based format into a temporary directory."""
    (tmp_path / "agreement.md").write_text(MARKDOWN, encoding="utf-8")
    (tmp_path / "agreement.html").write_text(HTML, encoding="utf-8")
    (tmp_path / "people.csv").write_text(CSV, encoding="utf-8")
    (tmp_path / "agreement.json").write_text(JSON_SOURCE, encoding="utf-8")
    (tmp_path / "notes.txt").write_text("First paragraph.\n\nSecond paragraph.\n", encoding="utf-8")
    (tmp_path / "agreement.docx").write_bytes(docx_bytes())
    (tmp_path / "agreement.pdf").write_bytes(pdf_bytes())
    return tmp_path
